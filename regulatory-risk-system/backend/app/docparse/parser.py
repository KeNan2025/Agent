"""
Document parser — produces structured chunks suitable for vector indexing.

Each chunk carries (text, page, section, metadata) so that the retrieval layer
can return precise citations like "2024年年报 P28".
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable


# Standard A-share annual-report section titles (per CSRC disclosure rules)
_SECTION_PATTERNS = [
    "重要事项", "公司基本情况", "经营情况讨论与分析",
    "财务报告", "审计报告", "股东变动",
    "董监高变动", "董事会报告", "监事会报告",
    "重要事项 关联交易", "重要事项",
    "财务报表附注",
]


@dataclass
class DocumentChunk:
    text: str
    chunk_id: str
    page: int = 0
    section: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    doc_id: str
    source: str
    chunks: list[DocumentChunk]
    metadata: dict[str, Any] = field(default_factory=dict)


# ─────────────────────────── Format dispatchers ───────────────────────────


def parse_file(path: str | Path, doc_id: str | None = None) -> Document:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"not found: {p}")
    suffix = p.suffix.lower()
    doc_id = doc_id or p.stem
    if suffix == ".pdf":
        return parse_pdf(p, doc_id)
    if suffix in {".html", ".htm"}:
        return parse_html(p, doc_id)
    if suffix in {".docx", ".doc"}:
        return parse_docx(p, doc_id)
    # Plain text fallback
    text = p.read_text(encoding="utf-8", errors="ignore")
    return Document(
        doc_id=doc_id, source=str(p),
        chunks=list(chunk_text(text, doc_id=doc_id)),
    )


# ─────────────────────────── PDF ───────────────────────────


def parse_pdf(path: Path, doc_id: str) -> Document:
    try:
        import fitz  # type: ignore  PyMuPDF
    except ImportError:
        # Fallback: read as binary, return single chunk acknowledging skip
        return Document(
            doc_id=doc_id, source=str(path),
            chunks=[DocumentChunk(
                text=f"[PyMuPDF not installed; PDF body not parsed: {path.name}]",
                chunk_id=f"{doc_id}_p0_c0", page=0, section="meta",
            )],
            metadata={"warning": "pymupdf_missing"},
        )

    doc = fitz.open(str(path))
    chunks: list[DocumentChunk] = []
    cur_section = ""
    for pno, page in enumerate(doc):
        text = page.get_text()
        if not text.strip():
            continue
        # Track section heading by simple pattern matching
        for ln in text.splitlines():
            ln_strip = ln.strip()
            if any(p in ln_strip for p in _SECTION_PATTERNS) and len(ln_strip) < 40:
                cur_section = ln_strip
                break
        for i, c in enumerate(chunk_text(text, doc_id=doc_id, page=pno + 1, section=cur_section)):
            c.chunk_id = f"{doc_id}_p{pno + 1}_c{i}"
            chunks.append(c)
    return Document(doc_id=doc_id, source=str(path), chunks=chunks)


# ─────────────────────────── HTML ───────────────────────────


def parse_html(path: Path, doc_id: str) -> Document:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        from bs4 import BeautifulSoup  # type: ignore
        soup = BeautifulSoup(raw, "html.parser")
        for s in soup(["script", "style"]):
            s.decompose()
        text = soup.get_text("\n")
    except ImportError:
        text = re.sub(r"<[^>]+>", "", raw)
    chunks = list(chunk_text(text, doc_id=doc_id))
    return Document(doc_id=doc_id, source=str(path), chunks=chunks)


# ─────────────────────────── DOCX ───────────────────────────


def parse_docx(path: Path, doc_id: str) -> Document:
    try:
        from docx import Document as DocxDoc  # type: ignore
    except ImportError:
        return Document(
            doc_id=doc_id, source=str(path),
            chunks=[DocumentChunk(
                text=f"[python-docx not installed; DOCX body not parsed: {path.name}]",
                chunk_id=f"{doc_id}_c0",
            )],
            metadata={"warning": "python_docx_missing"},
        )
    d = DocxDoc(str(path))
    text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
    chunks = list(chunk_text(text, doc_id=doc_id))
    return Document(doc_id=doc_id, source=str(path), chunks=chunks)


# ─────────────────────────── Chunking ───────────────────────────


def chunk_text(
    text: str, doc_id: str = "", page: int = 0, section: str = "",
    max_chars: int = 1024, overlap: int = 80,
) -> Iterable[DocumentChunk]:
    """Sliding-window character chunker with sentence-aware joining."""
    text = text.strip()
    if not text:
        return
    # Split into sentences first
    sentences = re.split(r"(?<=[。！？!?\.])\s+", text)
    buf: list[str] = []
    buf_len = 0
    idx = 0
    for sent in sentences:
        s = sent.strip()
        if not s:
            continue
        if buf_len + len(s) > max_chars and buf:
            yield DocumentChunk(
                text=" ".join(buf), chunk_id=f"{doc_id}_p{page}_c{idx}",
                page=page, section=section,
            )
            idx += 1
            # overlap = keep tail of buffer
            tail = " ".join(buf)[-overlap:] if overlap > 0 else ""
            buf = [tail, s] if tail else [s]
            buf_len = sum(len(x) for x in buf)
        else:
            buf.append(s)
            buf_len += len(s)
    if buf:
        yield DocumentChunk(
            text=" ".join(buf), chunk_id=f"{doc_id}_p{page}_c{idx}",
            page=page, section=section,
        )


def chunk_by_section(text: str, doc_id: str = "") -> list[DocumentChunk]:
    """Split by recognised CSRC section headings; falls back to sliding chunks."""
    pattern = re.compile(r"(?P<head>(?:" + "|".join(re.escape(s) for s in _SECTION_PATTERNS) + r"))")
    chunks: list[DocumentChunk] = []
    parts = pattern.split(text)
    if len(parts) <= 1:
        return list(chunk_text(text, doc_id=doc_id))
    idx = 0
    current_section = ""
    for i, p in enumerate(parts):
        if p in _SECTION_PATTERNS:
            current_section = p
            continue
        body = p.strip()
        if not body:
            continue
        for sub in chunk_text(body, doc_id=doc_id, section=current_section):
            sub.chunk_id = f"{doc_id}_s{idx}"
            chunks.append(sub)
            idx += 1
    return chunks
